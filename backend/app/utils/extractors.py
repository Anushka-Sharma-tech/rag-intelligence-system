import io
import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document
from typing import List, Dict

def extract_from_pdf(file_bytes: bytes, filename: str) -> List[Dict]:
    """Extract text from PDF page by page."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append({
                "text": text.strip(),
                "page": i + 1,
                "source": filename
            })
    return pages

def extract_from_docx(file_bytes: bytes, filename: str) -> List[Dict]:
    """Extract text from DOCX including tables."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            # Extract text from cells, filter out empty rows that just have " | " separators
            row_text = " | ".join(c.text.strip() for c in row.cells)
            if row_text.replace(" | ", "").strip(): 
                paragraphs.append(row_text)
                
    combined = "\n".join(paragraphs)
    
    # 🛑 NEW FIX: If the document contains absolutely no text, return an empty list so ingestion.py can catch it
    if not combined.strip():
        return []
        
    return [{"text": combined, "page": None, "source": filename}]

def extract_from_url(url: str) -> List[Dict]:
    """Extract readable text from a URL, stripping nav/footer/scripts."""
    headers = {"User-Agent": "Mozilla/5.0 (compatible; RAGBot/1.0)"}
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()
    soup = BeautifulSoup(response.content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    
    # 🛑 NEW FIX: Lowered character limit from 20 to 2 so we don't accidentally delete short bullet points or facts
    lines = [line.strip() for line in text.split("\n") if len(line.strip()) > 2]
    cleaned = "\n".join(lines)
    
    # 🛑 NEW FIX: If the URL had no readable text, return an empty list
    if not cleaned.strip():
        return []
        
    return [{"text": cleaned, "page": None, "source": url}]